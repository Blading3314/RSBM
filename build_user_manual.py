from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("UserManual_RSBM_RestaurantMGR.docx")


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="BFBFBF"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size in (("Title", 22), ("Heading 1", 15), ("Heading 2", 13)):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(6)


def add_page_number_header(section):
    section.different_first_page_header_footer = True
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.font.name = "Arial"
    run.font.size = Pt(10)
    add_field(p, "PAGE")


def add_cover(doc):
    for text, size, bold in [
        ("VANIER COLLEGE", 14, True),
        ("COMPUTER SCIENCE DEPARTMENT", 14, True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Arial"
        r.font.size = Pt(size)

    for _ in range(8):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("USER MANUAL")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)

    for _ in range(8):
        doc.add_paragraph()

    for line in [
        "JOEFFREY BRIONES, DANUSH SOOSAI",
        "420-331-VA : APPLICATION DEVELOPMENT 1 (DESKTOP)",
        "TEACHER: SYED NASEEM AFZAL",
        "MONDAY, MAY 11, 2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = "Arial"
        r.font.size = Pt(11)
        if line == "JOEFFREY BRIONES, DANUSH SOOSAI":
            r.bold = True

    doc.add_page_break()


def add_toc(doc):
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("TABLE OF CONTENTS")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(14)

    entries = [
        ("Introduction", 3),
        ("Interactive Features", 3),
        ("Graphical User Interface", 4),
        ("Basic UI", 4),
        ("Application Features", 5),
        ("Main Menu", 5),
        ("Reservations", 6),
        ("Exception Handling", 7),
        ("Tables", 8),
        ("Exception Handling", 9),
        ("Billing", 10),
        ("Internationalization", 11),
        ("Database and Output", 12),
        ("Previous Versions", 13),
    ]
    for name, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.add_run(name)
        p.add_run(" " + "." * max(10, 115 - len(name)))
        p.add_run(" {}".format(page))
    doc.add_page_break()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style="Heading {}".format(level))
    p.add_run(text.upper())


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(5)
        p.add_run(item)


def add_feature_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Area", True)
    set_cell_text(hdr[1], "User Action / Result", True)
    for cell in hdr:
        shade_cell(cell, "EDEDED")
        set_cell_border(cell)
    for left, right in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left, True)
        set_cell_text(cells[1], right)
        for cell in cells:
            set_cell_border(cell)
    doc.add_paragraph()


def add_intro(doc):
    add_heading(doc, "Introduction")
    doc.add_paragraph(
        "RSBM RestaurantMGR is a desktop restaurant management application designed for a Korean BBQ "
        "restaurant. The program helps staff manage reservations, track table availability, and prepare "
        "customer billing from a single Windows Forms interface."
    )
    doc.add_paragraph(
        "This user manual explains the main screens and day-to-day features of the application. It is "
        "written for restaurant staff who need a quick reference while seating guests, updating table "
        "statuses, and preparing bills."
    )
    add_heading(doc, "Interactive Features")
    add_bullets(
        doc,
        [
            "Language selector for English, French, and Spanish interface text.",
            "Navigation buttons for Reservations, Tables, and Billing.",
            "Reservation form for customer details, party size, date, time, table, and reservation status.",
            "Table manager for Korean BBQ grill seating such as two-seat, four-seat, family, and party tables.",
            "Billing screen for table information, AYCE pricing, extras, tip, total, and payment method.",
        ],
    )
    doc.add_page_break()


def add_gui(doc):
    add_heading(doc, "Graphical User Interface")
    add_feature_table(
        doc,
        [
            ("Header", "Displays the Restaurant Manager System title."),
            ("Sidebar Navigation", "Opens the Reservations, Tables, or Billing screen."),
            ("Language Selector", "Changes the displayed labels and messages to English, French, or Spanish."),
            ("Main Panel", "Displays the selected feature form without opening a separate application window."),
            ("Quit Button", "Closes the application when the user is finished."),
        ],
    )
    add_heading(doc, "Basic UI")
    doc.add_paragraph(
        "The application uses a simple dashboard layout with a dark header and sidebar. The user selects "
        "a feature from the left navigation area, and the selected form is loaded into the main content panel. "
        "This keeps the workflow consistent across the restaurant management screens."
    )
    doc.add_paragraph(
        "The language selector is located on the sidebar so staff can switch interface language without "
        "leaving the current workflow. When a language is selected, the main form and the currently opened "
        "child form refresh their displayed text."
    )
    doc.add_page_break()


def add_main_menu(doc):
    add_heading(doc, "Application Features")
    add_heading(doc, "Main Menu", 2)
    add_feature_table(
        doc,
        [
            ("Reservations", "Opens the reservation management screen."),
            ("Tables", "Opens the Korean BBQ table management screen."),
            ("Billing", "Opens the billing and payment screen."),
            ("Language", "Stores the selected language preference for the next time the program runs."),
        ],
    )
    doc.add_paragraph(
        "The main menu is intended to be the central starting point for restaurant staff. Each button opens "
        "a dedicated workspace for a common front-of-house task."
    )
    doc.add_page_break()


def add_reservations(doc):
    add_heading(doc, "Reservations")
    doc.add_paragraph(
        "The Reservations screen allows staff to enter customer information and reserve a table for a "
        "specific date and time. Staff can record the customer name, phone number, party size, reservation "
        "date, time slot, table, and current status."
    )
    add_feature_table(
        doc,
        [
            ("Customer Info", "Stores the customer name and phone number."),
            ("Reservation Details", "Stores party size, date, time, table, and status."),
            ("Status Options", "Tracks reservations as Confirmed, Pending, Seated, Completed, Cancelled, or No-show."),
            ("Search", "Provides a place to search reservations by name or ID."),
            ("Action Buttons", "Add, edit, delete, clear, or refresh reservation records."),
        ],
    )
    add_heading(doc, "Exception Handling", 2)
    add_bullets(
        doc,
        [
            "If no reservation is selected, the Edit and Delete actions display a message asking the user to select a reservation.",
            "Before deleting a reservation, the program asks for confirmation.",
            "When a reservation is added, updated, or deleted, a message box confirms the result.",
            "Database errors are shown through message boxes so the user knows that the action did not complete.",
        ],
    )
    doc.add_page_break()


def add_tables(doc):
    add_heading(doc, "Tables")
    doc.add_paragraph(
        "The Tables screen is used to manage Korean BBQ seating. Each table has a table number, capacity, "
        "Korean BBQ table type, and status. Staff can select a row in the grid and update the table's current "
        "status or seating type."
    )
    add_feature_table(
        doc,
        [
            ("Table Grid", "Displays table number, seat capacity, Korean BBQ table type, and status."),
            ("Selected Table", "Shows the table number selected from the grid."),
            ("Status", "Allows Available, Reserved, or Occupied."),
            ("Type", "Allows 2-seat BBQ grill, 4-seat BBQ grill, 6-seat family grill, or 8-seat party grill."),
            ("Update Table", "Saves the selected table's new capacity/type and status to the database."),
        ],
    )
    add_heading(doc, "Exception Handling", 2)
    add_bullets(
        doc,
        [
            "The user must select a table before pressing Update Table.",
            "The user must choose both a status and a table type.",
            "The table status is stored using valid database values only: Available, Reserved, and Occupied.",
            "If the table list cannot be loaded or saved, the application displays an error message.",
        ],
    )
    doc.add_page_break()


def add_billing(doc):
    add_heading(doc, "Billing")
    doc.add_paragraph(
        "The Billing screen helps staff prepare a bill for a table. It includes table information, AYCE "
        "pricing, optional extras, subtotal, tip, total, and payment method fields."
    )
    add_feature_table(
        doc,
        [
            ("Table Info", "Selects the customer table and displays related information."),
            ("Pricing", "Records the price per person and number of guests."),
            ("Extras", "Adds premium meat or drinks when selected."),
            ("Payment Method", "Tracks subtotal, tip, total, and payment method."),
            ("Generate Bill", "Produces the bill total for the current table order."),
        ],
    )
    add_heading(doc, "Exception Handling", 2)
    add_bullets(
        doc,
        [
            "Price and guest count should be entered before generating a bill.",
            "Optional extras are selected using check boxes to avoid typing errors.",
            "Payment totals should be reviewed before completing a transaction.",
        ],
    )
    doc.add_page_break()


def add_i18n_database(doc):
    add_heading(doc, "Internationalization")
    doc.add_paragraph(
        "The application supports English, French, and Spanish text through resource files. When the user "
        "changes the language selector, labels, buttons, messages, table status text, and Korean BBQ table "
        "type text are refreshed on the current screen."
    )
    add_feature_table(
        doc,
        [
            ("English", "Default application language."),
            ("French", "Displays translated labels, buttons, messages, and table values."),
            ("Spanish", "Displays translated labels, buttons, messages, and table values."),
            ("Saved Preference", "Stores the user's last selected language in application settings."),
        ],
    )
    add_heading(doc, "Database and Output")
    doc.add_paragraph(
        "The project uses a LocalDB database named RestaurantDB. The database stores customers, restaurant "
        "tables, reservations, and bills. The initialization script creates starter Korean BBQ tables with "
        "two-seat, four-seat, six-seat, and eight-seat grill capacities."
    )
    add_bullets(
        doc,
        [
            "Customers stores customer name and phone number.",
            "RestaurantTables stores table number, capacity, and status.",
            "Reservations connects customers to restaurant tables for a date and time.",
            "Bills stores subtotal, tax, total amount, payment method, payment status, and bill date.",
        ],
    )
    doc.add_page_break()


def add_versions(doc):
    add_heading(doc, "Previous Versions")
    add_heading(doc, "Version 1.0.0", 2)
    doc.add_paragraph(
        "Initial restaurant management interface with Reservations, Tables, and Billing navigation."
    )
    add_heading(doc, "Version 1.0.1", 2)
    doc.add_paragraph(
        "Added English, French, and Spanish internationalization support."
    )
    add_heading(doc, "Version 1.0.2", 2)
    doc.add_paragraph(
        "Completed Korean BBQ table management with localized table statuses and table types."
    )


def main():
    doc = Document()
    style_document(doc)
    add_page_number_header(doc.sections[0])
    add_cover(doc)
    add_toc(doc)
    add_intro(doc)
    add_gui(doc)
    add_main_menu(doc)
    add_reservations(doc)
    add_tables(doc)
    add_billing(doc)
    add_i18n_database(doc)
    add_versions(doc)
    doc.save(OUT)


if __name__ == "__main__":
    main()
